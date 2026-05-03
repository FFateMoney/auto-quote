from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from backend.common.config import PROJECT_ROOT
from backend.common.logging import append_run_log
from backend.quote.catalog import CatalogGateway
from backend.quote.form_ops import apply_manual_values
from backend.quote.kernel import Kernel
from backend.quote.llm.requester import QwenRequester
from backend.quote.models import FormRow, FormStageSnapshot, ResumeRequest, RunArtifacts, RunState, UploadedDocument
from backend.quote.plugins.registry import PluginRegistry
from backend.quote.quoter import Quoter
from backend.quote.run_store import RunStore
from backend.quote.stages import (
    DOCUMENT_EXTRACTED,
    EQUIPMENT_SELECTED_ENRICHED,
    EQUIPMENT_SELECTED_INITIAL,
    FINAL_QUOTED,
    STAGE_LABELS,
    STANDARD_ENRICHED,
    TEST_TYPE_MATCHED,
)
from backend.quote.standard.judge import StandardContextJudge
from backend.quote.standard.module import StandardRetrievalModule
from backend.quote.standard_enrich import progressive_enrich


logger = logging.getLogger(__name__)


def _copy(rows: list[FormRow]) -> list[FormRow]:
    return [r.model_copy(deep=True) for r in rows]


@dataclass(slots=True)
class QuoteOrchestrator:
    registry: PluginRegistry = field(default_factory=PluginRegistry)
    store: RunStore = field(default_factory=RunStore)
    catalog: CatalogGateway = field(default_factory=CatalogGateway)
    requester: QwenRequester = field(default_factory=QwenRequester)
    kernel: Kernel = field(init=False)
    quoter: Quoter = field(init=False)

    def __post_init__(self) -> None:
        judge = StandardContextJudge(requester=self.requester)
        retrieval = StandardRetrievalModule(judge=judge)
        self.kernel = Kernel(catalog=self.catalog, retrieval=retrieval)
        self.quoter = Quoter(catalog=self.catalog)

    # ------------------------------------------------------------------
    # Public
    # ------------------------------------------------------------------

    def run(self, *, run_id: str, uploaded_documents: list[UploadedDocument]) -> RunState:
        from backend.quote.settings import get_settings
        run_dir = get_settings().run_dir / run_id
        run_state_path = run_dir / "run_state.json"
        self._log(run_dir, "收到新运行请求，共 %s 个文件", len(uploaded_documents))

        state = RunState(
            run_id=run_id,
            overall_status="running",
            uploaded_documents=[d.model_copy(deep=True) for d in uploaded_documents],
            artifacts=RunArtifacts(
                run_state_path=str(run_state_path),
                uploaded_dir=str(run_dir / "uploaded"),
            ),
            next_action="系统正在处理文档",
        )
        self._save(run_state_path, state)

        try:
            documents, preprocess_notes = self._preprocess(state, run_dir)
            self._log(run_dir, "文档解析完成，共 %s 份", len(documents))

            extraction = self.requester.extract_form(documents, run_dir=run_dir)
            rows = extraction.items
            notes = list(preprocess_notes)
            if extraction.summary:
                notes.append(f"模型摘要：{extraction.summary}")
            self._upsert(state, DOCUMENT_EXTRACTED, rows, notes)
            self._log(run_dir, "文件抽取完成，共 %s 行", len(rows))
            self._save(run_state_path, state)

            rows, notes = self.kernel.match_test_types(rows)
            self._upsert(state, TEST_TYPE_MATCHED, rows, notes)
            self._log(run_dir, "试验类型匹配完成，共 %s 行", len(rows))
            self._save(run_state_path, state)

            rows, planning_notes = self.quoter.plan_standard_fields(rows)
            rows, notes = self.quoter.select_equipment(rows)
            self._upsert(state, EQUIPMENT_SELECTED_INITIAL, rows, [*planning_notes, *notes])
            self._log(run_dir, "设备筛选完成，共 %s 行", len(rows))
            self._save(run_state_path, state)

            rows = self._standard_stage(state, rows, run_dir)
            self._log(run_dir, "标准补充完成，共 %s 行", len(rows))
            self._save(run_state_path, state)

            rows, notes = self.quoter.select_equipment(rows)
            rows, repeat_notes = self.quoter.assign_repeat_counts(rows)
            self._upsert(state, EQUIPMENT_SELECTED_ENRICHED, rows, ["标准补充后重新筛选设备", *notes, *repeat_notes])
            self._log(run_dir, "标准补充后设备筛选完成，共 %s 行", len(rows))
            self._save(run_state_path, state)

            rows, notes, status = self.quoter.price(rows)
            self._upsert(state, FINAL_QUOTED, rows, notes)
            state.final_form_items = _copy(rows)
            state.current_stage = FINAL_QUOTED
            state.overall_status = status
            state.next_action = self._next_action(status)
            self._log(run_dir, "最终报价完成，状态=%s，行数=%s", status, len(rows))
            self._save(run_state_path, state)
        except Exception as exc:
            state.overall_status = "failed"
            state.errors.append(str(exc))
            state.next_action = "检查错误信息后重新上传文档"
            self._log(run_dir, "运行失败: %s", str(exc), level=logging.ERROR)
            self._save(run_state_path, state)

        return state

    def resume(self, *, run_id: str, request: ResumeRequest) -> RunState:
        from backend.quote.settings import get_settings
        run_dir = get_settings().run_dir / run_id
        run_state_path = run_dir / "run_state.json"
        state = self.store.load(run_state_path)
        self._log(run_dir, "收到人工补录继续报价请求: row_id=%s", request.row_id)

        test_type_changed = "canonical_test_type" in request.field_values
        base_rows = state.final_form_items or self._stage_rows(state, FINAL_QUOTED)
        if not test_type_changed:
            base_rows = self._stage_rows(state, STANDARD_ENRICHED) or base_rows
        rows = apply_manual_values(base_rows, request.row_id, request.field_values)
        try:
            if test_type_changed:
                rows, notes = self.kernel.match_test_types(rows)
                self._upsert(state, TEST_TYPE_MATCHED, rows, ["用户补录后重新匹配试验类型", *notes])

                rows, planning_notes = self.quoter.plan_standard_fields(rows)
                rows, notes = self.quoter.select_equipment(rows)
                self._upsert(state, EQUIPMENT_SELECTED_INITIAL, rows, ["用户补录后重新筛选设备", *planning_notes, *notes])

                rows = self._standard_stage(state, rows, run_dir)
            else:
                rows = self._clear_equipment_state(rows)
                self._upsert(
                    state,
                    STANDARD_ENRICHED,
                    rows,
                    ["用户补录未修改标准试验类型，复用已有标准补充结果，跳过标准文档检索"],
                )

            rows, notes = self.quoter.select_equipment(rows)
            rows, repeat_notes = self.quoter.assign_repeat_counts(rows)
            self._upsert(state, EQUIPMENT_SELECTED_ENRICHED, rows, ["标准补充后重新筛选设备", *notes, *repeat_notes])

            rows, notes, status = self.quoter.price(rows)
            self._upsert(state, FINAL_QUOTED, rows, ["用户补录后重新报价", *notes])
            state.final_form_items = _copy(rows)
            state.current_stage = FINAL_QUOTED
            state.overall_status = status
            state.next_action = self._next_action(status)
            self._save(run_state_path, state)
        except Exception as exc:
            state.overall_status = "failed"
            state.errors.append(str(exc))
            state.next_action = "人工补录后重新报价失败，请检查输入"
            self._log(run_dir, "人工补录后重新报价失败: %s", str(exc), level=logging.ERROR)
            self._save(run_state_path, state)

        return state

    def load_run(self, run_id: str) -> RunState:
        from backend.quote.settings import get_settings
        return self.store.load(get_settings().run_dir / run_id / "run_state.json")

    # ------------------------------------------------------------------
    # Private
    # ------------------------------------------------------------------

    def _preprocess(self, state: RunState, run_dir: Path) -> tuple[list, list[str]]:
        documents, notes = [], []
        for uploaded in state.uploaded_documents:
            path = Path(uploaded.local_path or uploaded.stored_path)
            plugin = self.registry.resolve(path)
            self._log(run_dir, "文档路由: %s -> %s", uploaded.file_name, plugin.plugin_id)
            normalized = plugin.preprocess(path, {"run_dir": run_dir})
            uploaded.source_kind = normalized.source_kind
            uploaded.status = "preprocessed"
            uploaded.notes = f"plugin={plugin.plugin_id}"
            documents.append(normalized)
            notes.append(f"{uploaded.file_name}: routed to {plugin.plugin_id}")
        return documents, notes

    def _standard_stage(self, state: RunState, rows: list[FormRow], run_dir: Path) -> list[FormRow]:
        rows = self._clear_standard_discovery_state(rows)
        target_by_row = {
            row.row_id: list(row.planned_standard_fields)
            for row in rows
            if row.standard_codes and row.planned_standard_fields
        }
        if not target_by_row:
            rows = self._clear_equipment_state(rows)
            self._upsert(state, STANDARD_ENRICHED, rows, ["无标准补充模板字段或缺少标准号，跳过标准补充"])
            return rows

        rows, evidence_notes = self.kernel.resolve_standard_evidences(rows, target_fields_by_row=target_by_row, run_dir=run_dir)
        notes = list(evidence_notes)

        candidate_rows = [r for r in rows if target_by_row.get(r.row_id) and r.standard_evidences]
        if candidate_rows:
            discovery = self.requester.discover_standard_fields(
                candidate_rows,
                target_fields_by_row=target_by_row,
                supported_fields=self.quoter.supported_standard_fields(),
                run_dir=run_dir,
            )
            if discovery.summary:
                notes.append(f"字段发现摘要：{discovery.summary}")
            rows = self._apply_standard_discovery(rows, discovery.items)
            notes.extend(self._discovery_notes(rows, discovery.items))

            fill_targets_by_row = {
                row.row_id: [
                    field_name
                    for field_name in row.discovered_standard_fields
                    if not _has_value(getattr(row, field_name, None))
                ]
                for row in rows
                if row.row_id in target_by_row
            }
            fill_targets_by_row = {row_id: fields for row_id, fields in fill_targets_by_row.items() if fields}

            if fill_targets_by_row:
                rows, enrich_notes = progressive_enrich(
                    rows, target_fields_by_row=fill_targets_by_row, requester=self.requester, run_dir=run_dir
                )
                notes.extend(enrich_notes)
            else:
                notes.append("字段发现未产出新的待补值字段，跳过标准补值")
        elif any(target_by_row.values()):
            notes.append("存在标准补充目标字段，但未命中可用标准证据，按当前信息继续报价")
        else:
            notes.append("未命中有效标准证据，按当前信息继续报价")

        rows = self._clear_equipment_state(rows)
        self._upsert(state, STANDARD_ENRICHED, rows, notes)
        return rows

    def _apply_standard_discovery(
        self,
        rows: list[FormRow],
        items,
    ) -> list[FormRow]:
        discovery_by_row = {item.row_id: item for item in items}
        updated: list[FormRow] = []
        for row in rows:
            copy = row.model_copy(deep=True)
            copy.discovered_standard_fields = []
            copy.extra_standard_requirements = []
            result = discovery_by_row.get(row.row_id)
            if result is None:
                updated.append(copy)
                continue
            copy.discovered_standard_fields = list(result.discovered_standard_fields)
            copy.extra_standard_requirements = [item.model_copy(deep=True) for item in result.extra_standard_requirements]
            updated.append(copy)
        return updated

    def _clear_standard_discovery_state(self, rows: list[FormRow]) -> list[FormRow]:
        updated: list[FormRow] = []
        for row in rows:
            copy = row.model_copy(deep=True)
            copy.discovered_standard_fields = []
            copy.extra_standard_requirements = []
            copy.standard_evidences = []
            copy.standard_match_notes = []
            updated.append(copy)
        return updated

    def _clear_equipment_state(self, rows: list[FormRow]) -> list[FormRow]:
        updated: list[FormRow] = []
        for row in rows:
            copy = row.model_copy(deep=True)
            copy.candidate_equipment_ids = []
            copy.candidate_equipment_profiles = []
            copy.selected_equipment_id = ""
            copy.rejected_equipment = []
            copy.missing_fields = []
            copy.blocking_reason = ""
            updated.append(copy)
        return updated

    def _discovery_notes(self, rows: list[FormRow], items) -> list[str]:
        rows_by_id = {row.row_id: row for row in rows}
        notes: list[str] = []
        for item in items:
            row = rows_by_id.get(item.row_id)
            label = (row.raw_test_type or row.canonical_test_type or item.row_id) if row else item.row_id
            if item.discovered_standard_fields:
                notes.append(f"{label}: 发现标准字段 {', '.join(item.discovered_standard_fields)}")
            else:
                notes.append(f"{label}: 未发现新的系统支持字段")
            if item.extra_standard_requirements:
                notes.append(f"{label}: 记录额外标准要求 {len(item.extra_standard_requirements)} 条")
        return notes

    def _upsert(self, state: RunState, stage_id: str, rows: list[FormRow], notes: list[str]) -> None:
        snapshot = FormStageSnapshot(
            stage_id=stage_id,
            label=STAGE_LABELS[stage_id],
            items=_copy(rows),
            notes=[n for n in notes if n],
        )
        for i, stage in enumerate(state.form_stages):
            if stage.stage_id == stage_id:
                state.form_stages[i] = snapshot
                state.current_stage = stage_id
                return
        state.form_stages.append(snapshot)
        state.current_stage = stage_id

    def _stage_rows(self, state: RunState, stage_id: str) -> list[FormRow]:
        for stage in state.form_stages:
            if stage.stage_id == stage_id:
                return _copy(stage.items)
        return []

    def export_docx(self, run_id: str) -> Path:
        from backend.quote.settings import get_settings
        state = self.load_run(run_id)
        settings = get_settings()
        template_path = PROJECT_ROOT / "doc" / "quote_tep.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        run_dir = settings.run_dir / run_id
        output_path = run_dir / f"报价单_{run_id}.docx"

        doc = Document(template_path)

        # 1. 填充表格 (Table 1: 序号 | 测试项目 | 备注 | 基本金(元) | 单价(元) | 试验总量 | 合计（元）)
        if len(doc.tables) > 1:
            table = doc.tables[1]
            # 查找数据起始行（序号为1的行）
            start_row_idx = -1
            for i, row in enumerate(table.rows):
                if row.cells[0].text.strip() == "1":
                    start_row_idx = i
                    break
            
            if start_row_idx != -1:
                items = state.final_form_items
                
                # 识别模板中现有的数字序号行（数据行）
                data_row_indices = []
                total_row_idx = -1
                for i in range(start_row_idx, len(table.rows)):
                    cell_text = table.rows[i].cells[0].text.strip()
                    if cell_text.isdigit():
                        data_row_indices.append(i)
                    elif "总计" in cell_text:
                        total_row_idx = i
                        break
                
                # 如果实际项目数多于模板预设行数，在总计行之前插入新行
                if len(items) > len(data_row_indices):
                    # 如果没有找到总计行，就在末尾加，否则在总计行之前加
                    insert_before_idx = total_row_idx if total_row_idx != -1 else len(table.rows)
                    for _ in range(len(items) - len(data_row_indices)):
                        # 注意：python-docx 的 add_row 总是加在末尾
                        # 如果需要插入，逻辑会复杂些，这里简单处理：先加行，后面填充时按序号来
                        table.add_row()
                    
                    # 重新刷新行索引
                    data_row_indices = []
                    total_row_idx = -1
                    for i in range(start_row_idx, len(table.rows)):
                        cell_text = table.rows[i].cells[0].text.strip()
                        if cell_text.isdigit() or cell_text == "": # 包含新加的空行
                            data_row_indices.append(i)
                        elif "总计" in cell_text:
                            total_row_idx = i
                            break

                # 填充数据或清空多余模板行
                total_amount = 0.0
                for i, row_idx in enumerate(data_row_indices):
                    row = table.rows[row_idx]
                    if i < len(items):
                        item = items[i]
                        row.cells[0].text = str(i + 1)
                        row.cells[1].text = item.canonical_test_type or item.raw_test_type or ""
                        row.cells[2].text = "" # 备注留空
                        row.cells[3].text = f"{item.base_fee:g}" if item.base_fee is not None else "0"
                        row.cells[4].text = f"{item.unit_price:g}" if item.unit_price is not None else "0"
                        row.cells[5].text = f"{item.pricing_quantity:g}" if item.pricing_quantity is not None else "0"
                        row.cells[6].text = f"{item.total_price:g}" if item.total_price is not None else "0"
                        total_amount += (item.total_price or 0.0)
                    else:
                        # 清空多余的模板行内容
                        for cell in row.cells:
                            cell.text = ""

                # 更新总计行
                if total_row_idx != -1:
                    table.rows[total_row_idx].cells[6].text = f"{total_amount:g}"

        # 2. 替换文本占位符 (日期和注)
        now = datetime.now()
        date_str = now.strftime("%Y年%m月%d日")
        
        for para in doc.paragraphs:
            if "20xx年x月x日" in para.text:
                para.text = para.text.replace("20xx年x月x日", date_str)
        
        # 处理表格内的占位符（注：部分通常在最后一行的大单元格里）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "x" in cell.text or "xxxx" in cell.text:
                        # 注1: 样品名 (留空)
                        new_text = cell.text.replace("xxxx", "")
                        # 注3: 项目数量
                        new_text = new_text.replace("x个测试项目", f"{len(state.final_form_items)}个测试项目")
                        cell.text = new_text

        doc.save(output_path)
        
        # 更新状态中的产物列表
        if str(output_path.name) not in state.artifacts.exported_files:
            state.artifacts.exported_files.append(output_path.name)
            self._save(run_dir / "run_state.json", state)
            
        return output_path

    def _save(self, path: Path, state: RunState) -> None:
        self.store.save(path, state)

    def _log(self, run_dir: Path, message: str, *args: object, level: int = logging.INFO) -> None:
        text = message % args if args else message
        logger.log(level, "[run=%s] %s", run_dir.name, text)
        append_run_log(run_dir, text)

    def _next_action(self, status: str) -> str:
        if status == "completed":
            return "查看最终报价表或下载产物"
        if status == "waiting_manual_input":
            return "补齐表格中的缺失字段后继续报价"
        return "检查系统错误后重试"


def _has_value(value: object) -> bool:
    return value not in (None, "", [])
